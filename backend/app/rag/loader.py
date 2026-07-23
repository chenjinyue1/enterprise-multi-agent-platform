
# 3. 编写 RAG Loader 模块
"""
RAG 文档加载模块
负责加载报告模板文件，进行预处理和分块

小白理解：
Loader就像"文件扫描仪+剪刀"。它做两件事：
1. 扫描：读取各种格式的文件（PDF、Word、Markdown、Excel）
2. 裁剪：把长文档切成小段（Chunk），每段300-500字

为什么要切？因为大模型一次能看的文字有限（上下文窗口），
而且切小块后，检索更精准——就像书分成章节后更容易找到相关内容。
"""

import os
import re
from typing import List, Optional
from pathlib import Path

from langchain_core.documents import Document
from langchain.text_splitter import (
    RecursiveCharacterTextSplitter,
    MarkdownHeaderTextSplitter
)


class ReportTemplateLoader:
    """
    报告模板加载器
    
    支持的文件格式：
    ----------------
    - .md    Markdown报告模板（最推荐，结构清晰）
    - .txt   纯文本模板
    - .docx  Word文档（需要python-docx）
    - .pdf   PDF报告（需要PyPDF2或pdfplumber）
    
    企业场景：
    ---------
    公司历史上有几百份报告，格式各异：
    - 销售部用Excel
    - 市场部用PPT转PDF
    - 管理层用Word
    - 技术部用Markdown
    
    Loader统一处理所有格式，提取纯文本+元数据，然后分块入库。
    """
    
    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        supported_extensions: Optional[List[str]] = None
    ):
        """
        初始化加载器
        
        Args:
            chunk_size: 每个文本块的大小（token数），500是报告场景的黄金分割点
            chunk_overlap: 块之间重叠的字数，防止关键信息被切分在两块边界
            supported_extensions: 支持的文件扩展名列表
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.supported_extensions = supported_extensions or [".md", ".txt", ".docx", ".pdf"]
        
        # 通用文本分割器（按字符递归分割，优先在段落/句子边界切）
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\\n\\n", "\\n", "。", "；", " ", ""],
            length_function=len
        )
        
        # Markdown专用分割器（按标题层级切分，保留结构）
        self.markdown_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=[
                ("#", "header_1"),
                ("##", "header_2"),
                ("###", "header_3")
            ]
        )
    
    def load_directory(
        self,
        directory: str,
        report_type: Optional[str] = None,
        department: Optional[str] = None
    ) -> List[Document]:
        """
        批量加载目录下的所有模板文件
        
        Args:
            directory: 模板文件目录（如 ./data/templates/）
            report_type: 统一标记这批文档的报告类型
            department: 统一标记这批文档的所属部门
            
        Returns:
            Document对象列表，每个Document包含：
            - page_content: 文本内容
            - metadata: {source, report_type, department, ...}
            
        使用示例：
        -------
        loader = ReportTemplateLoader()
        docs = loader.load_directory(
            directory="./data/templates/quarterly_reports",
            report_type="quarterly",
            department="sales"
        )
        # docs 可以直接传给 EmbeddingService.add_documents()
        """
        documents = []
        directory_path = Path(directory)
        
        if not directory_path.exists():
            print(f"⚠️ 目录不存在: {directory}")
            return documents
        
        # 遍历目录下所有支持的文件
        for ext in self.supported_extensions:
            for file_path in directory_path.rglob(f"*{ext}"):
                try:
                    docs = self.load_file(file_path)
                    
                    # 统一添加元数据
                    for doc in docs:
                        doc.metadata.update({
                            "source": str(file_path),
                            "report_type": report_type or self._infer_report_type(file_path),
                            "department": department or "unknown",
                            "file_ext": ext,
                            "load_time": str(datetime.now())
                        })
                    
                    documents.extend(docs)
                    print(f"✅ 已加载: {file_path}")
                    
                except Exception as e:
                    print(f"❌ 加载失败: {file_path} - {e}")
        
        return documents
    
    def load_file(self, file_path: str) -> List[Document]:
        """
        加载单个文件
        
        根据文件扩展名自动选择对应的解析方法
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        
        if ext == ".md":
            return self._load_markdown(file_path)
        elif ext == ".txt":
            return self._load_text(file_path)
        elif ext == ".docx":
            return self._load_docx(file_path)
        elif ext == ".pdf":
            return self._load_pdf(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
    
    def _load_markdown(self, file_path: str) -> List[Document]:
        """加载Markdown文件（保留标题层级结构）"""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        
        # 尝试按标题分割（保留报告结构）
        try:
            docs = self.markdown_splitter.split_text(content)
            # 补充文件名到metadata
            for doc in docs:
                doc.metadata["source"] = file_path
            return docs
        except Exception:
            # 如果Markdown结构不规范，回退到通用分割
            return self.text_splitter.create_documents(
                texts=[content],
                metadatas=[{"source": file_path}]
            )
    
    def _load_text(self, file_path: str) -> List[Document]:
        """加载纯文本文件"""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        
        return self.text_splitter.create_documents(
            texts=[content],
            metadatas=[{"source": file_path}]
        )
    
    def _load_docx(self, file_path: str) -> List[Document]:
        """加载Word文档"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            full_text = []
            
            # 提取段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        full_text.append(row_text)
            
            content = "\\n".join(full_text)
            
            return self.text_splitter.create_documents(
                texts=[content],
                metadatas=[{"source": file_path}]
            )
        except ImportError:
            raise ImportError("加载.docx需要安装python-docx: pip install python-docx")
    
    def _load_pdf(self, file_path: str) -> List[Document]:
        """加载PDF文件"""
        try:
            import pdfplumber
            
            full_text = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        full_text.append(text)
            
            content = "\\n\\n".join(full_text)
            
            return self.text_splitter.create_documents(
                texts=[content],
                metadatas=[{"source": file_path}]
            )
        except ImportError:
            raise ImportError("加载PDF需要安装pdfplumber: pip install pdfplumber")
    
    def _infer_report_type(self, file_path: Path) -> str:
        """
        从文件路径推断报告类型
        
        例如：
        - .../quarterly/... → quarterly
        - .../monthly/... → monthly
        - .../annual/... → annual
        """
        path_str = str(file_path).lower()
        
        if "quarter" in path_str or "q1" in path_str or "q2" in path_str:
            return "quarterly"
        elif "month" in path_str or "monthly" in path_str:
            return "monthly"
        elif "annual" in path_str or "year" in path_str:
            return "annual"
        elif "special" in path_str or "event" in path_str:
            return "special"
        else:
            return "general"
    
    def create_sample_templates(self, output_dir: str = "./data/templates") -> None:
        """
        创建示例报告模板（用于演示和测试）
        
        企业场景：新项目启动时，先创建几个示例模板入库，
        后续Report Agent就有"历史经验"可以参考了。
        """
        os.makedirs(output_dir, exist_ok=True)
        
        templates = {
            "quarterly_sales_template.md": """# 季度销售分析报告

## 一、执行摘要

本季度销售总体表现{{summary}}。核心指标如下：
- 总销售额：{{total_sales}}万元
- 同比增长：{{yoy_growth}}%
- 环比增长：{{mom_growth}}%

## 二、各品类销售分析

{{category_analysis}}

## 三、区域销售分布

{{region_analysis}}

## 四、关键发现与建议

{{insights}}

## 五、下季度展望

{{outlook}}
""",
            "monthly_operation_template.md": """# 月度运营报告

## 一、本月核心数据

- GMV：{{gmv}}万元
- 订单量：{{orders}}单
- 客单价：{{aov}}元
- 转化率：{{conversion}}%

## 二、流量分析

{{traffic_analysis}}

## 三、用户行为分析

{{user_behavior}}

## 四、问题与改进

{{issues_and_improvements}}
""",
            "annual_summary_template.md": """# 年度总结报告

## 一、年度业绩回顾

{{annual_performance}}

## 二、重点项目复盘

{{project_review}}

## 三、团队成长

{{team_growth}}

## 四、明年规划

{{next_year_plan}}
"""
        }
        
        for filename, content in templates.items():
            filepath = os.path.join(output_dir, filename)
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(content)
            print(f"✅ 创建示例模板: {filepath}")


# 便捷函数
def load_and_index_templates(
    template_dir: str = "./data/templates",
    report_type: Optional[str] = None
) -> int:
    """
    一键加载并索引所有模板
    
    使用示例：
    -------
    count = load_and_index_templates("./data/templates", "quarterly")
    print(f"成功索引 {count} 个文档块")
    """
    from .embeddings import get_embedding_service
    
    loader = ReportTemplateLoader()
    
    # 加载文档
    docs = loader.load_directory(template_dir, report_type=report_type)
    
    if not docs:
        print("⚠️ 未找到可加载的模板文件")
        return 0
    
    # 向量化入库
    embedding_service = get_embedding_service()
    ids = embedding_service.add_documents(docs)
    
    print(f"✅ 成功索引 {len(ids)} 个文档块")
    return len(ids)


print("✅ loader.py 编写完成")
